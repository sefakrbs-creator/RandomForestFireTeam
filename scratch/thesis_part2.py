"""Part 2: Adds Methodology, Results, Discussion, Conclusion to the thesis doc."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

path = r"c:\Users\sefak\OneDrive\Documents\sefakarabas\OneDrive\Desktop\Dataset-2\Shrubwise_Thesis.docx"
doc = Document(path)

# ---- CHAPTER 4: METHODOLOGY ----
doc.add_heading('Chapter 4: Methodology', level=1)
doc.add_paragraph(
    'This chapter describes the complete processing pipeline, from raw field data to operational '
    'shrub maps. The methodology is organized into five stages: (1) ground truth generation from '
    'TLS, (2) patch extraction and tiling, (3) SAM-based annotation refinement, (4) feature '
    'engineering and model training, and (5) inference and deployment.'
)

doc.add_heading('4.1 Ground Truth Generation from TLS', level=2)
doc.add_paragraph(
    'Shrub locations from field survey CSVs are projected onto the NAIP coordinate grid using '
    'rigid-body transformation matrices. Each surveyed shrub is represented as a circular mask '
    'centered on its projected coordinates with a radius proportional to its measured canopy '
    'diameter. These initial circular masks serve as seed annotations for subsequent SAM refinement.'
)
doc.add_paragraph(
    'This step is methodologically important because it establishes the spatial link between '
    'field-verified ecological observations and the image domain. Without this georeferencing, '
    'no supervised learning is possible. The use of TLS rather than GPS-only coordinates ensures '
    'sub-meter positional accuracy, which is critical when working with 1-meter resolution imagery '
    'where a single pixel displacement can shift a label from shrub to bare ground.'
)

doc.add_heading('4.2 Patch Extraction and Tiling', level=2)
doc.add_paragraph(
    'Full-site rasters are tiled into 128x128 pixel patches, each containing 5 channels: NAIP '
    'Red, Green, Blue, NIR, and Canopy Height. Patches are stored as NumPy arrays (.npy) and '
    'partitioned into training (677 patches) and validation (202 patches) splits stratified by '
    'site to prevent spatial autocorrelation between splits.'
)
doc.add_paragraph(
    'The 128x128 pixel patch size was chosen as a compromise between spatial context (128 meters '
    'at 1m resolution provides sufficient neighborhood for multi-scale texture computation) and '
    'computational tractability for SAM processing. Site-stratified splitting ensures that the '
    'model is evaluated on its ability to generalize across landscapes rather than memorizing '
    'site-specific patterns, which is essential for a model intended for deployment at new locations.'
)

doc.add_heading('4.3 SAM-Based Annotation Refinement', level=2)
doc.add_paragraph(
    'The Segment Anything Model (SAM ViT-H, 632M parameters) is applied in two passes to '
    'convert the initial circular annotations into ecologically validated, pixel-accurate labels.'
)

doc.add_heading('4.3.1 Pass 1: Prompted Shrub Refinement', level=3)
doc.add_paragraph(
    'For each circular shrub annotation, the centroid coordinates are used as a foreground point '
    'prompt to SAM. The model generates a precise segmentation mask that conforms to the actual '
    'shrub boundary visible in the NAIP imagery. The refined mask replaces the original circle '
    'only if its mean CHM value falls within the valid shrub height range (1.0-4.0 m). Patches '
    'are upscaled 4x (128 to 512 pixels) before SAM processing to improve small-object detection.'
)
doc.add_paragraph(
    'This height-based validation is a critical quality control mechanism. SAM is a general-purpose '
    'segmentation model with no ecological knowledge; it may segment visually salient objects '
    'that are not actually shrubs (e.g., shadows, soil patches). The CHM filter ensures that only '
    'masks corresponding to vegetation in the 1-4 meter height range are accepted as shrub labels, '
    'grounding the annotation process in physical measurement.'
)

doc.add_heading('4.3.2 Pass 2: Hard Negative Mining via AMG', level=3)
doc.add_paragraph(
    'SAM\'s Automatic Mask Generator (AMG) is run on every patch to discover unlabeled objects. '
    'Each generated mask is classified based on its mean CHM value: masks with CHM > 4.0 m are '
    'labeled as "tree" (hard negative), while masks with CHM < 1.0 m are labeled as "rock/ground" '
    '(hard negative). Masks overlapping existing shrub annotations (IoU > 0.3) are discarded to '
    'prevent label conflicts. Area filters (4-3000 pixels) remove noise and background.'
)
doc.add_paragraph(
    'Hard negative mining addresses a fundamental challenge in binary classification: the model '
    'must learn not only what shrubs look like, but also what they do not look like. Without '
    'explicit tree and rock examples, the classifier tends to over-predict shrubs in forested '
    'areas (confusing tree canopy with shrub canopy) and rocky terrain (confusing textured rock '
    'surfaces with shrub patches). The AMG-based approach is fully automated and scales to '
    'thousands of patches without human intervention.'
)

doc.add_heading('4.4 Feature Engineering', level=2)
doc.add_paragraph(
    'A total of 55 features are computed per pixel from the 5-channel input patches. These are '
    'organized into four groups:'
)
doc.add_paragraph('Raw spectral bands: Red, Green, Blue, NIR (4 features)', style='List Bullet')
doc.add_paragraph(
    'Vegetation indices: NDVI, SAVI, GNDVI, EVI2, brightness, greenness ratio (6 features)',
    style='List Bullet'
)
doc.add_paragraph(
    'Canopy height derivatives: raw height, canopy_in_shrub_range (binary flag), '
    'canopy_shrub_clipped (3 features)', style='List Bullet'
)
doc.add_paragraph(
    'Multi-scale texture: mean and standard deviation at 3x3, 5x5, and 7x7 pixel windows '
    'computed over 7 base bands (NDVI, NIR, Red, Green, SAVI, brightness, canopy height), '
    'yielding 42 texture features', style='List Bullet'
)
doc.add_paragraph(
    'Texture features are computed using uniform filtering with reflect padding at patch boundaries '
    'to avoid edge artifacts. The multi-scale approach captures spatial context at different '
    'neighborhood sizes: 3x3 (3m) resolves individual shrub crowns, 5x5 (5m) captures shrub '
    'clumps, and 7x7 (7m) captures landscape-level vegetation patterns.'
)

doc.add_heading('4.5 Model Training', level=2)
doc.add_paragraph(
    'Training follows a two-pass strategy using XGBoost (Chen & Guestrin, 2016):'
)

doc.add_heading('4.5.1 Pass 1: Feature Selection', level=3)
doc.add_paragraph(
    'An initial XGBoost model is trained on all 55 features with conservative hyperparameters '
    '(learning_rate=0.05, max_depth=6, n_estimators=500). Features are ranked by split-based '
    'importance, and the top 20 are selected for the second pass. This reduces dimensionality '
    'by 64%, removing redundant and noisy features while retaining those with the strongest '
    'discriminative power.'
)

doc.add_heading('4.5.2 Pass 2: Multi-Objective Hyperparameter Optimization', level=3)
doc.add_paragraph(
    'An 80-trial Optuna search using the NSGA-II multi-objective sampler simultaneously maximizes '
    'three objectives: F1 score, precision, and recall. This multi-objective formulation is '
    'critical for ecological applications where the cost of false positives (wasted field '
    'verification effort) and false negatives (missed shrubs affecting fuel models) must be '
    'explicitly balanced. The Pareto-optimal trial with the highest F1 is selected, and the '
    'final model is trained with 2,000 estimators and early stopping (patience=50 rounds).'
)
doc.add_paragraph(
    'Features are scaled using RobustScaler (median/IQR normalization) which is resistant to '
    'outliers common in remote sensing data. The binary flag canopy_in_shrub_range is excluded '
    'from scaling to preserve its discrete nature.'
)
doc.add_page_break()

# ---- CHAPTER 5: RESULTS ----
doc.add_heading('Chapter 5: Results', level=1)

doc.add_heading('5.1 Model Performance', level=2)
doc.add_paragraph(
    'The V12 XGBoost classifier was evaluated on the held-out validation set (202 patches from '
    'all six sites). The following table summarizes the classification performance:'
)
t = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (m, v) in enumerate([
    ('F1 Score', '0.823'), ('Precision', '0.753'),
    ('Recall', '0.907'), ('AUC-ROC', '0.772'),
    ('Optimal Threshold', '0.544')
]):
    t.rows[i].cells[0].text = m
    t.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph(
    'The high recall (0.907) indicates that the model successfully identifies over 90% of true '
    'shrub pixels, which is desirable for fuel mapping applications where missed shrubs could '
    'lead to underestimation of fire risk. The precision of 0.753 means approximately 25% of '
    'predicted shrub pixels are false positives, primarily occurring at the edges of tree canopies '
    'where height transitions from tree to shrub range.'
)

doc.add_heading('5.2 Feature Importance Analysis', level=2)
doc.add_paragraph(
    'The final model\'s feature importance ranking reveals that canopy height texture features '
    'dominate the classification decision:'
)
t = doc.add_table(rows=11, cols=3, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.rows[0].cells[0].text = 'Rank'; t.rows[0].cells[1].text = 'Feature'; t.rows[0].cells[2].text = 'Importance'
for c in range(3): t.rows[0].cells[c].paragraphs[0].runs[0].bold = True
for r, (rk, ft, im) in enumerate([
    ('1','canopy_height_7x7mean','0.234'), ('2','canopy_height_5x5mean','0.140'),
    ('3','canopy_height_3x3mean','0.067'), ('4','brightness_7x7std','0.057'),
    ('5','canopy_height_7x7std','0.055'), ('6','naip_nir_7x7std','0.044'),
    ('7','ndvi_7x7mean','0.040'), ('8','greenness_ratio','0.038'),
    ('9','brightness_5x5std','0.037'), ('10','canopy_in_shrub_range','0.036'),
], 1):
    t.rows[r].cells[0].text = rk; t.rows[r].cells[1].text = ft; t.rows[r].cells[2].text = im

doc.add_paragraph()
doc.add_paragraph(
    'The top three features are all canopy height means at different spatial scales, collectively '
    'accounting for 44% of total importance. This strongly supports the hypothesis that structural '
    'information (height) is more discriminative than spectral information for shrub classification. '
    'Notably, raw spectral band values (naip_red, naip_green, naip_blue, naip_nir) were eliminated '
    'during feature selection, appearing only in texture-derived forms. This confirms that spatial '
    'context (local mean and variability) is more informative than individual pixel values.'
)

doc.add_heading('5.3 Per-Site Prediction Results', level=2)
t = doc.add_table(rows=7, cols=5, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Site','Raster Size','Shrub (ha)','p50','p95']):
    t.rows[0].cells[i].text = h; t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for r, row in enumerate([
    ('Calaveras Big Trees','1537x4107','182.5','0.394','0.851'),
    ('DL Bliss','1124x578','40.8','0.667','0.915'),
    ('Independence Lake','1161x1760','151.3','0.733','0.921'),
    ('Pacific Union','2716x1087','219.1','0.735','0.889'),
    ('Sedgwick','1815x589','103.9','0.872','0.945'),
    ('Shaver Lake','1078x531','35.6','0.655','0.909'),
], 1):
    for c, v in enumerate(row): t.rows[r].cells[c].text = v

doc.add_paragraph()
doc.add_paragraph(
    'Total mapped shrub area across all six sites is approximately 733 hectares. The variation in '
    'median probability (p50) across sites reflects ecological differences: Sedgwick (coastal sage '
    'scrub, p50=0.872) has the highest confidence because its dominant shrub species are spectrally '
    'and structurally distinct, while Calaveras (mixed conifer, p50=0.394) shows lower confidence '
    'due to the complexity of the understory beneath the conifer canopy.'
)

doc.add_heading('5.4 Comparison with Alternative Approaches', level=2)
t = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Approach','Version','F1','Key Limitation']):
    t.rows[0].cells[i].text = h; t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for r, row in enumerate([
    ('Detectron2 Mask R-CNN','v5-v7','0.000','Instance seg failed on small objects'),
    ('U-Net Semantic Seg','v8-v9','Moderate','Patch boundary artifacts, heavy compute'),
    ('LightGBM + CHM labels','v10','Good','Noisy labels (CHM threshold only)'),
    ('XGBoost + SAM labels','v12','0.823','Best balance of all metrics'),
], 1):
    for c, v in enumerate(row): t.rows[r].cells[c].text = v
doc.add_page_break()

# ---- CHAPTER 6: DISCUSSION ----
doc.add_heading('Chapter 6: Discussion', level=1)

doc.add_heading('6.1 Importance of Structural Features', level=2)
doc.add_paragraph(
    'The dominance of canopy height features in the model importance ranking has important '
    'implications for shrub mapping methodology. It suggests that purely spectral approaches '
    '(using only NAIP bands and vegetation indices) are fundamentally limited for distinguishing '
    'shrubs from trees, because both vegetation types can exhibit similar NDVI and brightness '
    'values. The critical discriminative information lies in the height domain: shrubs occupy '
    'a well-defined structural niche (1-4 m) that separates them from both ground-level features '
    '(<1 m) and tree canopy (>4 m). As global canopy height products (e.g., Meta\'s 1m CHM) '
    'become increasingly available and accurate, this structural approach will become applicable '
    'to regions where no local LiDAR data exists.'
)

doc.add_heading('6.2 Role of SAM in Label Quality', level=2)
doc.add_paragraph(
    'The transition from CHM-threshold labels (V10, LightGBM) to SAM-refined polygon labels (V12, '
    'XGBoost) produced a significant improvement in model quality, despite using the same underlying '
    'spectral and structural features. This demonstrates that label quality is at least as important '
    'as model architecture for ecological classification tasks. SAM provides two orthogonal '
    'improvements: (1) precise positive boundaries that reduce label noise at shrub edges, and '
    '(2) systematically mined hard negatives that explicitly teach the model to reject trees and '
    'rocks. The automated nature of SAM-based annotation makes it scalable to datasets far larger '
    'than what manual digitization could achieve.'
)

doc.add_heading('6.3 Why Gradient Boosting Outperformed Deep Learning', level=2)
doc.add_paragraph(
    'The failure of Detectron2 Mask R-CNN (AP=0.000) and the moderate performance of U-Net '
    'highlight a common pitfall in applying deep learning to ecological remote sensing. These '
    'architectures require thousands of labeled instances to learn effective representations, but '
    'the Shrubwise dataset contains only 879 annotated patches. In contrast, XGBoost operates on '
    'hand-engineered features that encode domain knowledge (vegetation indices, height thresholds, '
    'multi-scale texture) directly, allowing the model to achieve strong performance from limited '
    'data. Additionally, XGBoost provides interpretable feature importances that validate the '
    'ecological reasoning behind the model, an important requirement for scientific applications '
    'where black-box predictions are insufficient.'
)

doc.add_heading('6.4 Limitations', level=2)
doc.add_paragraph(
    'Several limitations should be noted. First, the validation set shares the same six sites as '
    'the training set (though patches are spatially separated), so generalization to entirely '
    'new landscapes has not been formally evaluated. Second, the 0.544 classification threshold '
    'was optimized for the training sites and may require adjustment for different ecosystems. '
    'Third, the pipeline assumes temporal alignment between NAIP imagery and CHM data, which may '
    'not hold in areas with rapid land cover change.'
)
doc.add_page_break()

# ---- CHAPTER 7: CONCLUSION ----
doc.add_heading('Chapter 7: Conclusion and Future Work', level=1)

doc.add_heading('7.1 Contributions', level=2)
doc.add_paragraph('This dissertation makes the following contributions:')
doc.add_paragraph(
    'A novel annotation pipeline combining SAM with CHM-based validation that automatically '
    'converts point-level field observations into pixel-accurate shrub labels and mines hard '
    'negative examples.', style='List Bullet'
)
doc.add_paragraph(
    'An XGBoost pixel classifier achieving F1=0.823 for shrub detection using only NAIP imagery '
    'and canopy height, demonstrating that engineered features with gradient boosting can outperform '
    'deep learning for ecological classification with limited training data.', style='List Bullet'
)
doc.add_paragraph(
    'Empirical evidence that canopy height texture features account for approximately 50% of '
    'classification importance, establishing structural information as the primary discriminative '
    'signal for shrub mapping.', style='List Bullet'
)
doc.add_paragraph(
    'An autonomous prediction agent that enables operational shrub mapping for arbitrary areas '
    'of interest without manual data preparation.', style='List Bullet'
)

doc.add_heading('7.2 Future Work', level=2)
doc.add_paragraph(
    'Several directions for future research emerge from this work. First, formal cross-site '
    'validation using a leave-one-site-out protocol would provide a more rigorous assessment of '
    'generalization. Second, temporal analysis using multi-year NAIP acquisitions could enable '
    'shrub change detection and growth monitoring. Third, integration with GEDI waveform data '
    'could improve height estimation in areas where the Meta CHM product has limited accuracy. '
    'Finally, the SAM annotation framework could be extended to multi-class vegetation mapping '
    '(distinguishing shrub species) by incorporating species-specific height and spectral profiles.'
)

doc.save(path)
print(f"Thesis saved to: {path}")
