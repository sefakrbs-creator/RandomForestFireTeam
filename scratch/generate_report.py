"""
generate_report.py - Creates a Word document report for the Shrubwise repository.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

# --- Title Page ---
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Shrubwise: Automated Shrub Detection\nfrom Aerial Imagery')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Technical Report')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('XGBoost Pixel Classification Pipeline (V12)\nwith SAM-Refined Annotations')
run.font.size = Pt(13)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('April 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ===== 1. EXECUTIVE SUMMARY =====
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This repository implements an end-to-end machine learning pipeline for mapping shrub cover '
    'across California field sites using 1-meter resolution NAIP aerial imagery and LiDAR-derived '
    'Canopy Height Models (CHM). The system combines Meta\'s Segment Anything Model (SAM) for '
    'precise ground truth annotation with an XGBoost pixel classifier for scalable inference.'
)
doc.add_paragraph(
    'The final production model (V12) achieves F1 = 0.823, Precision = 0.753, and Recall = 0.907 '
    'on the held-out validation set. An autonomous agent can generate shrub maps for any new area '
    'of interest in the continental US given only a GeoJSON boundary file.'
)

# Key metrics table
doc.add_heading('Key Performance Metrics', level=2)
t = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for r, (metric, val) in enumerate([
    ('F1 Score', '0.823'), ('Precision', '0.753'),
    ('Recall', '0.907'), ('AUC-ROC', '0.772'),
    ('Optimal Threshold', '0.544')
]):
    t.rows[r].cells[0].text = metric
    t.rows[r].cells[1].text = val

# ===== 2. PROBLEM STATEMENT =====
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'Accurate mapping of shrub cover is critical for wildfire fuel load estimation, habitat '
    'assessment, and ecological monitoring. Traditional field surveys are slow, expensive, and '
    'limited in spatial coverage. Remote sensing offers a scalable alternative, but distinguishing '
    'shrubs from trees, rocks, and bare ground in aerial imagery remains challenging due to their '
    'small size, variable shape, and spectral similarity to surrounding vegetation.'
)
doc.add_paragraph(
    'This project automates shrub detection by: (1) training a pixel-level classifier on '
    'expert-verified labels refined by SAM, (2) predicting shrub probability maps at 1m resolution '
    'across six study sites, and (3) deploying an autonomous agent that can map shrubs anywhere '
    'in the continental US.'
)

# ===== 3. STUDY SITES =====
doc.add_heading('3. Study Sites', level=1)
doc.add_paragraph(
    'The pipeline was developed and validated across six California field sites spanning diverse '
    'vegetation types from coastal chaparral to subalpine forest:'
)
t = doc.add_table(rows=7, cols=4, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Site', 'Code', 'Scan Dates', 'Ecosystem']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
sites = [
    ('Calaveras Big Trees', 'CATCU', 'Oct 2024, May-Jun 2025', 'Mixed conifer forest'),
    ('DL Bliss State Park', 'CAAEU', 'Multiple', 'Sierra Nevada montane'),
    ('Independence Lake', '-', 'Jul-Aug 2024, Aug 2025', 'Subalpine lake basin'),
    ('Pacific Union College', '-', 'Nov 2024', 'N. California woodland'),
    ('Sedgwick Reserve', 'CASBC', 'Sep-Nov 2024, Jun-Jul 2025', 'Coastal sage scrub'),
    ('Shaver Lake', 'CAFKU', 'Jul 2024', 'Sierra foothills'),
]
for r, row_data in enumerate(sites, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph(
    'Each site contains Terrestrial LiDAR Scans (TLS), 4-band NAIP imagery at 1m resolution, '
    'Canopy Height Models from Meta\'s global dataset, and field-verified shrub survey lists. '
    'All rasters use EPSG:6350 (NAD83(2011) / Conus Albers) as the common coordinate system.'
)

# ===== 4. PIPELINE ARCHITECTURE =====
doc.add_heading('4. Pipeline Architecture', level=1)
doc.add_paragraph(
    'The pipeline follows a seven-step workflow. Each step builds on the previous one and addresses '
    'a specific technical challenge. Below, each step is described along with a rationale for why '
    'it is essential to the overall system.'
)

# --- Step 1 ---
doc.add_heading('Step 1: TLS to Mask Rasters (process_shrub_masks.py)', level=2)
doc.add_paragraph(
    'This step projects Terrestrial LiDAR Scan (TLS) point clouds onto the NAIP coordinate grid '
    'to create georeferenced shrub mask rasters. Each surveyed shrub location from the field CSV '
    'is converted into a circular mask polygon at its known coordinates.'
)
p = doc.add_paragraph()
run = p.add_run('Why is this step important? ')
run.bold = True
p.add_run(
    'Without spatially accurate ground truth labels aligned to the imagery, no supervised model '
    'can learn what a shrub looks like. TLS provides centimeter-level accuracy for shrub locations, '
    'but these 3D coordinates must be projected into the 2D image space before they can serve as '
    'training labels. This step bridges the gap between field measurements and image-based learning.'
)

# --- Step 2 ---
doc.add_heading('Step 2: Patch Tiling (128x128 px, 5-channel .npy)', level=2)
doc.add_paragraph(
    'Full-site NAIP rasters (often thousands of pixels wide) are too large to process as single '
    'images. This step tiles them into 128x128 pixel patches, each containing 5 channels: '
    'Red, Green, Blue, NIR (from NAIP) and Canopy Height (from CHM). The patches are stored as '
    'NumPy arrays (.npy) and split into 677 training and 202 validation patches.'
)
p = doc.add_paragraph()
run = p.add_run('Why is this step important? ')
run.bold = True
p.add_run(
    'Tiling serves three purposes: (1) it makes the data manageable for SAM processing, which '
    'expects bounded images; (2) it enables stratified train/val splitting by site to prevent '
    'spatial autocorrelation bias; and (3) it allows per-patch texture feature computation with '
    'proper boundary handling (reflect padding) instead of grid-edge artifacts.'
)

# --- Step 3 ---
doc.add_heading('Step 3: Label Studio Export (export_patches_for_labeling.py)', level=2)
doc.add_paragraph(
    'Patches are exported as 512x512 PNGs (4x upscaled for visibility) with existing circular '
    'annotations overlaid. A Label Studio import JSON is generated for manual review.'
)
p = doc.add_paragraph()
run = p.add_run('Why is this step important? ')
run.bold = True
p.add_run(
    'Automated labels inevitably contain errors. This step provides a human-in-the-loop quality '
    'control checkpoint where domain experts can verify, correct, or reject annotations before '
    'they are used for training. Visual review catches systematic issues such as misaligned TLS '
    'projections or incorrect coordinate transforms that would silently corrupt the model.'
)

# --- Step 4 ---
doc.add_heading('Step 4: SAM Annotation Refinement (sam_annotate.py)', level=2)
doc.add_paragraph(
    'Meta\'s Segment Anything Model (SAM ViT-H) refines annotations in two passes:'
)
doc.add_paragraph(
    'Pass 1 (Prompted SAM): Each circular shrub annotation\'s centroid is used as a foreground '
    'point prompt. SAM generates a precise mask conforming to the actual shrub boundary in the '
    'imagery. The refined mask replaces the circle only if its mean CHM falls within [1.0, 4.0] m.',
    style='List Bullet'
)
doc.add_paragraph(
    'Pass 2 (Automatic Mask Generation): SAM\'s AMG runs on every patch to discover unlabeled '
    'objects. Masks with CHM > 4m are labeled as "tree" hard negatives; CHM < 1m as "rock" hard '
    'negatives. These teach the model what shrubs are NOT.',
    style='List Bullet'
)
p = doc.add_paragraph()
run = p.add_run('Why is this step important? ')
run.bold = True
p.add_run(
    'Circular annotations are geometrically crude approximations of irregularly shaped shrubs. '
    'SAM refinement converts these into pixel-accurate boundaries that follow the actual shrub '
    'edges, dramatically improving label quality. Furthermore, the hard negative mining is crucial: '
    'without explicit tree and rock examples, the model struggles to distinguish shrubs from other '
    'vegetation or terrain features with similar spectral signatures. The CHM-validated filtering '
    'ensures that SAM\'s masks are ecologically meaningful rather than purely visual.'
)

# Annotation stats table
doc.add_heading('Annotation Statistics', level=3)
t = doc.add_table(rows=8, cols=2, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.rows[0].cells[0].text = 'Site'
t.rows[0].cells[1].text = 'Patches'
t.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t.rows[0].cells[1].paragraphs[0].runs[0].bold = True
for r, (s, n) in enumerate([
    ('Calaveras Big Trees', '463'), ('Sedgwick', '151'),
    ('Independence Lake', '132'), ('Pacific Union', '53'),
    ('DL Bliss', '40'), ('Shaver Lake', '40'), ('Total', '879')
], 1):
    t.rows[r].cells[0].text = s
    t.rows[r].cells[1].text = n

# --- Step 5 ---
doc.add_heading('Step 5: XGBoost Training (train_shrub_v12.py)', level=2)
doc.add_paragraph(
    'The production model is a binary per-pixel XGBoost classifier trained in two passes:'
)
doc.add_paragraph(
    'Pass 1 (Feature Selection): 55 candidate features are computed per pixel, including raw '
    'spectral bands, vegetation indices (NDVI, SAVI, GNDVI, EVI2), brightness, greenness ratio, '
    'canopy height derivatives, and multi-scale texture statistics (mean/std at 3x3, 5x5, 7x7 '
    'windows). XGBoost importance ranking selects the top 20 features.',
    style='List Bullet'
)
doc.add_paragraph(
    'Pass 2 (Hyperparameter Optimization): An 80-trial Optuna NSGA-II multi-objective search '
    'simultaneously maximizes F1, Precision, and Recall. The Pareto-optimal trial with the '
    'highest F1 is selected. The final model uses 2,000 estimators with early stopping.',
    style='List Bullet'
)
p = doc.add_paragraph()
run = p.add_run('Why is this step important? ')
run.bold = True
p.add_run(
    'Feature selection prevents overfitting and reduces inference time by eliminating 35 redundant '
    'features. The multi-objective optimization is critical because ecological mapping requires '
    'balancing precision (avoiding false positives that waste field verification effort) against '
    'recall (not missing actual shrubs). Single-objective optimization on accuracy alone would '
    'produce a model biased toward the majority class (non-shrub pixels).'
)

# Feature importance table
doc.add_heading('Top 10 Features by Importance', level=3)
t = doc.add_table(rows=11, cols=3, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.rows[0].cells[0].text = 'Rank'
t.rows[0].cells[1].text = 'Feature'
t.rows[0].cells[2].text = 'Importance'
for c in range(3):
    t.rows[0].cells[c].paragraphs[0].runs[0].bold = True
feats = [
    ('1', 'canopy_height_7x7mean', '0.234'),
    ('2', 'canopy_height_5x5mean', '0.140'),
    ('3', 'canopy_height_3x3mean', '0.067'),
    ('4', 'brightness_7x7std', '0.057'),
    ('5', 'canopy_height_7x7std', '0.055'),
    ('6', 'naip_nir_7x7std', '0.044'),
    ('7', 'ndvi_7x7mean', '0.040'),
    ('8', 'greenness_ratio', '0.038'),
    ('9', 'brightness_5x5std', '0.037'),
    ('10', 'canopy_in_shrub_range', '0.036'),
]
for r, (rank, feat, imp) in enumerate(feats, 1):
    t.rows[r].cells[0].text = rank
    t.rows[r].cells[1].text = feat
    t.rows[r].cells[2].text = imp

doc.add_paragraph()
doc.add_paragraph(
    'Key finding: Canopy height features dominate the top ranks (positions 1, 2, 3, 5), accounting '
    'for approximately 50% of cumulative importance. This confirms that LiDAR-derived structural '
    'information is the strongest predictor of shrub presence. Multi-scale texture windows '
    'consistently outperform raw spectral values, suggesting that spatial context around each pixel '
    'is more informative than the pixel value alone.'
)

# --- Step 6 ---
doc.add_heading('Step 6: Batch Inference (predict_v12_all_sites.py)', level=2)
doc.add_paragraph(
    'The trained model is applied to full-site NAIP + CHM rasters for all six study sites. '
    'For each site, the script computes all 20 features at every pixel, runs XGBoost inference '
    'in 500,000-pixel chunks, and produces both a float32 probability raster (0-1) and a uint8 '
    'binary mask at the optimized threshold.'
)
p = doc.add_paragraph()
run = p.add_run('Why is this step important? ')
run.bold = True
p.add_run(
    'Batch inference validates that the model generalizes beyond the small training patches to '
    'full landscape-scale rasters. It also produces the deliverable outputs: georeferenced '
    'probability maps that can be loaded directly into GIS software (QGIS, ArcGIS) for analysis, '
    'overlay with other ecological datasets, and field planning.'
)

# Results table
doc.add_heading('Prediction Results (threshold = 0.544)', level=3)
t = doc.add_table(rows=7, cols=5, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Site', 'Raster Size', 'Shrub Area (ha)', 'p50', 'p95']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
results = [
    ('Calaveras Big Trees', '1537 x 4107', '182.5', '0.394', '0.851'),
    ('DL Bliss', '1124 x 578', '40.8', '0.667', '0.915'),
    ('Independence Lake', '1161 x 1760', '151.3', '0.733', '0.921'),
    ('Pacific Union', '2716 x 1087', '219.1', '0.735', '0.889'),
    ('Sedgwick', '1815 x 589', '103.9', '0.872', '0.945'),
    ('Shaver Lake', '1078 x 531', '35.6', '0.655', '0.909'),
]
for r, row_data in enumerate(results, 1):
    for c, val in enumerate(row_data):
        t.rows[r].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph('Total mapped shrub area: approximately 733 hectares across 6 sites.')

# --- Step 7 ---
doc.add_heading('Step 7: Autonomous Agent (autonomous_shrub_agent.py)', level=2)
doc.add_paragraph(
    'The autonomous agent enables inference on any new area of interest by automating data '
    'acquisition. Given a GeoJSON boundary file, it: (1) initializes Google Earth Engine, '
    '(2) downloads NAIP imagery from Microsoft Planetary Computer, (3) downloads CHM from GEE, '
    '(4) aligns both rasters to a common 1m grid, and (5) runs V12 XGBoost inference.'
)
p = doc.add_paragraph()
run = p.add_run('Why is this step important? ')
run.bold = True
p.add_run(
    'Without the agent, using the model on a new site requires manually downloading and aligning '
    'NAIP and CHM rasters - a process that involves navigating multiple data portals, handling '
    'coordinate reprojection, and ensuring pixel-level alignment. The agent reduces this to a '
    'single command, making the pipeline accessible to ecologists and land managers who may not '
    'have geospatial programming expertise. It transforms the model from a research tool into '
    'a practical field-ready application.'
)

# ===== 5. MODEL DESIGN DECISIONS =====
doc.add_heading('5. Design Decisions', level=1)

doc.add_heading('Why XGBoost over Deep Learning?', level=2)
doc.add_paragraph(
    'The project explored multiple approaches before settling on XGBoost:'
)
t = doc.add_table(rows=5, cols=3, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.rows[0].cells[0].text = 'Approach'
t.rows[0].cells[1].text = 'Result'
t.rows[0].cells[2].text = 'Issue'
for c in range(3):
    t.rows[0].cells[c].paragraphs[0].runs[0].bold = True
approaches = [
    ('Detectron2 Mask R-CNN (v5-v7)', 'AP = 0.000', 'Instance seg failed; class imbalance'),
    ('U-Net CNN (v8-v9)', 'Moderate', 'Patch boundary artifacts, heavy compute'),
    ('LightGBM (v10)', 'Good', 'Labels were noisy (CHM-threshold only)'),
    ('XGBoost + SAM (v12)', 'F1 = 0.823', 'Best accuracy, speed, interpretability'),
]
for r, (app, res, issue) in enumerate(approaches, 1):
    t.rows[r].cells[0].text = app
    t.rows[r].cells[1].text = res
    t.rows[r].cells[2].text = issue

doc.add_paragraph()
doc.add_paragraph(
    'XGBoost on hand-engineered features outperformed deep learning for this task because: '
    '(1) the training dataset (879 patches) is too small for CNNs to generalize; '
    '(2) pixel-level features with multi-scale texture capture the relevant spatial context; '
    '(3) XGBoost provides interpretable feature importances that validate ecological reasoning; '
    'and (4) inference is fast enough for landscape-scale application without GPU hardware.'
)

doc.add_heading('Why SAM for Annotation?', level=2)
doc.add_paragraph(
    'Traditional circular annotations from field survey coordinates are geometrically crude. '
    'SAM provides pixel-accurate boundaries that conform to actual shrub edges, automated hard '
    'negative mining (trees and rocks are systematically identified), and scalability - processing '
    '879 patches automatically with CHM-based validation.'
)

doc.add_heading('Why No Terrain Features?', level=2)
doc.add_paragraph(
    'V12 intentionally excludes elevation, slope, and aspect. Canopy height alone captures the '
    'structural signal needed for shrub vs. tree/rock discrimination. Removing terrain features '
    'improves generalizability to new sites where DEM alignment may differ and reduces the number '
    'of data dependencies for the autonomous agent.'
)

# ===== 6. DEPENDENCIES =====
doc.add_heading('6. Dependencies', level=1)
t = doc.add_table(rows=6, cols=2, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.rows[0].cells[0].text = 'Category'
t.rows[0].cells[1].text = 'Packages'
t.rows[0].cells[0].paragraphs[0].runs[0].bold = True
t.rows[0].cells[1].paragraphs[0].runs[0].bold = True
deps = [
    ('Core ML', 'numpy, pandas, scipy, scikit-learn, xgboost, optuna, joblib'),
    ('Geospatial', 'rasterio, geopandas, shapely'),
    ('Remote Sensing', 'earthengine-api, planetary-computer, pystac-client'),
    ('Computer Vision', 'opencv-python, segment-anything, torch, torchvision, Pillow'),
    ('Visualization', 'matplotlib'),
]
for r, (cat, pkgs) in enumerate(deps, 1):
    t.rows[r].cells[0].text = cat
    t.rows[r].cells[1].text = pkgs

# ===== 7. HOW TO REPRODUCE =====
doc.add_heading('7. How to Reproduce', level=1)

doc.add_heading('Train the model', level=2)
doc.add_paragraph('source venv_linux/bin/activate', style='No Spacing')
doc.add_paragraph('python train_shrub_v12.py', style='No Spacing')

doc.add_heading('Run inference on known sites', level=2)
doc.add_paragraph('python predict_v12_all_sites.py', style='No Spacing')
doc.add_paragraph('python predict_v12_all_sites.py --sites DL_Bliss', style='No Spacing')
doc.add_paragraph('python predict_v12_all_sites.py --threshold 0.9', style='No Spacing')

doc.add_heading('Predict a new area of interest', level=2)
doc.add_paragraph(
    'python Model_Prediction/autonomous_shrub_agent.py my_area.geojson '
    '--output results/ --year 2022 --threshold 0.9',
    style='No Spacing'
)

# --- Save ---
out_path = Path(r"c:\Users\sefak\OneDrive\Documents\sefakarabas\OneDrive\Desktop\Dataset-2\Shrubwise_Report.docx")
doc.save(str(out_path))
print(f"Report saved to: {out_path}")
