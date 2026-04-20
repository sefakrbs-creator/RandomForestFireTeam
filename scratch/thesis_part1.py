"""Part 1: Create thesis doc with Intro, Lit Review, Study Area. Saves as pickle for Part 2."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pickle

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.5

for lv in range(1,4):
    hs = doc.styles[f'Heading {lv}']
    hs.font.color.rgb = RGBColor(0,0,0)
    hs.font.name = 'Times New Roman'
    hs.font.bold = True

# ---- TITLE PAGE ----
for _ in range(4): doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('AUTOMATED SHRUB COVER MAPPING FROM\nMULTI-SOURCE REMOTE SENSING DATA\nUSING MACHINE LEARNING')
r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'
doc.add_paragraph()
tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp2.add_run('A Thesis Presented to the Faculty\nin Partial Fulfillment of the Requirements\nfor the Degree of Doctor of Philosophy')
r.font.size = Pt(14); r.font.name = 'Times New Roman'
doc.add_paragraph()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp3.add_run('April 2026')
r.font.size = Pt(14)
doc.add_page_break()

# ---- ABSTRACT ----
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'Shrub cover mapping is essential for wildfire fuel load estimation, carbon stock accounting, '
    'and biodiversity monitoring in Mediterranean and semi-arid ecosystems. Traditional field-based '
    'methods are spatially limited and labor-intensive, while existing remote sensing approaches '
    'struggle to distinguish shrubs from trees and bare ground due to spectral similarity and '
    'fine spatial scale. This dissertation presents a novel automated pipeline that integrates '
    'Terrestrial LiDAR Scanning (TLS), the Segment Anything Model (SAM) for annotation refinement, '
    'and gradient-boosted decision trees (XGBoost) for pixel-level shrub classification from '
    '1-meter NAIP aerial imagery and LiDAR-derived Canopy Height Models (CHM).'
)
doc.add_paragraph(
    'The proposed methodology was developed and validated across six ecologically diverse field '
    'sites in California, spanning coastal chaparral, mixed conifer forest, and subalpine '
    'ecosystems. A two-pass annotation strategy using SAM converts imprecise circular ground '
    'truth markers into pixel-accurate shrub boundaries while simultaneously mining hard negative '
    'examples (trees and rocks) through height-based filtering. The resulting XGBoost classifier '
    '(V12) achieves an F1 score of 0.823, precision of 0.753, and recall of 0.907 on a held-out '
    'validation set using only 20 engineered features selected from an initial pool of 55 candidates.'
)
doc.add_paragraph(
    'A key finding is that canopy height texture features dominate model importance (approximately '
    '50% of cumulative weight), confirming that structural information from LiDAR is more '
    'discriminative than spectral signatures alone. The pipeline is operationalized through an '
    'autonomous agent that automatically acquires NAIP imagery and canopy height data for any '
    'user-specified area of interest, enabling scalable shrub mapping across the continental '
    'United States without manual data preparation.'
)
doc.add_page_break()

# ---- CHAPTER 1: INTRODUCTION ----
doc.add_heading('Chapter 1: Introduction', level=1)

doc.add_heading('1.1 Background and Motivation', level=2)
doc.add_paragraph(
    'Shrublands constitute a significant fraction of global terrestrial vegetation, covering '
    'approximately 45 million square kilometers worldwide (Staver et al., 2011). In California '
    'alone, chaparral and coastal sage scrub ecosystems represent over 8% of the state\'s land '
    'area and serve critical ecological functions including soil stabilization, carbon '
    'sequestration, wildlife habitat provision, and watershed protection (Keeley & Davis, 2007). '
    'Despite their ecological importance, shrubs remain among the most poorly mapped vegetation '
    'types due to their intermediate structural characteristics: they are too small for reliable '
    'detection by moderate-resolution satellite sensors (e.g., Landsat at 30m) yet too numerous '
    'for exhaustive field inventory.'
)
doc.add_paragraph(
    'The increasing frequency and severity of wildfires in western North America has created an '
    'urgent operational need for accurate shrub cover maps. Shrubs constitute a primary component '
    'of surface fuels that drive fire spread rates and flame lengths (Rothermel, 1972). Current '
    'fuel mapping products, such as LANDFIRE, rely on 30-meter resolution classifications that '
    'cannot capture the fine-scale heterogeneity of shrub distributions critical for fire behavior '
    'modeling (Rollins, 2009). High-resolution (1-meter) shrub maps would enable more accurate '
    'fire risk assessments, targeted fuel treatment planning, and post-fire recovery monitoring.'
)

doc.add_heading('1.2 Research Gap', level=2)
doc.add_paragraph(
    'Several limitations persist in existing approaches to shrub mapping from remote sensing data. '
    'First, most studies rely on spectral information alone, which is insufficient to distinguish '
    'shrubs from trees when both exhibit similar greenness indices (NDVI, SAVI). Second, the '
    'generation of training labels for supervised classification remains a major bottleneck: '
    'field surveys provide point-level locations but not pixel-accurate boundaries, requiring '
    'manual digitization that is prohibitively time-consuming at scale. Third, deep learning '
    'approaches (e.g., semantic segmentation with U-Net or instance segmentation with Mask R-CNN) '
    'require large labeled datasets that are rarely available for ecological applications, '
    'particularly for vegetation classes that lack clear spectral distinctiveness.'
)
doc.add_paragraph(
    'This dissertation addresses these gaps by proposing a pipeline that: (1) leverages LiDAR-derived '
    'canopy height as the primary discriminative feature, exploiting the structural difference '
    'between shrubs (1-4 m) and trees (>4 m); (2) uses the Segment Anything Model (SAM) to '
    'automatically refine imprecise field annotations into pixel-accurate labels while mining '
    'hard negative examples; and (3) employs gradient-boosted decision trees that are robust to '
    'limited training data and provide interpretable feature importance rankings.'
)

doc.add_heading('1.3 Research Objectives', level=2)
doc.add_paragraph('The specific objectives of this research are:', style='List Bullet')
doc.add_paragraph(
    'To develop an automated annotation refinement pipeline using SAM that converts point-level '
    'field observations into pixel-accurate shrub boundary polygons validated by canopy height.',
    style='List Bullet'
)
doc.add_paragraph(
    'To design and train an XGBoost pixel classifier using multi-source features (spectral, '
    'textural, and structural) that achieves high accuracy in distinguishing shrubs from trees, '
    'rocks, and bare ground.',
    style='List Bullet'
)
doc.add_paragraph(
    'To evaluate the relative importance of spectral vs. structural features for shrub detection '
    'across ecologically diverse California landscapes.',
    style='List Bullet'
)
doc.add_paragraph(
    'To operationalize the pipeline through an autonomous agent enabling shrub mapping for '
    'arbitrary areas of interest without manual data preparation.',
    style='List Bullet'
)

doc.add_heading('1.4 Thesis Organization', level=2)
doc.add_paragraph(
    'This thesis is organized as follows. Chapter 2 reviews the literature on shrub mapping, '
    'remote sensing classification methods, and foundation models for annotation. Chapter 3 '
    'describes the six study sites and their ecological characteristics. Chapter 4 presents '
    'the methodology including data acquisition, SAM annotation refinement, feature engineering, '
    'and model training. Chapter 5 reports the experimental results and model performance. '
    'Chapter 6 discusses the findings, design decisions, and limitations. Chapter 7 concludes '
    'with contributions and directions for future research.'
)
doc.add_page_break()

# ---- CHAPTER 2: LITERATURE REVIEW ----
doc.add_heading('Chapter 2: Literature Review', level=1)

doc.add_heading('2.1 Remote Sensing of Shrub Vegetation', level=2)
doc.add_paragraph(
    'Remote sensing of shrub vegetation has evolved from coarse-resolution satellite classification '
    'to fine-scale aerial imagery analysis. Early approaches used Landsat-derived vegetation indices '
    '(NDVI, EVI) to map broad shrubland extent at 30-meter resolution (Xian et al., 2009). While '
    'effective for regional-scale mapping, these methods cannot resolve individual shrub patches '
    'or distinguish shrubs from the understory of open woodlands. The National Agriculture Imagery '
    'Program (NAIP) provides 1-meter, 4-band (R, G, B, NIR) imagery across the continental US, '
    'enabling sub-canopy vegetation mapping that was previously impossible from space.'
)
doc.add_paragraph(
    'The integration of LiDAR-derived height information has proven transformative for vegetation '
    'classification. Airborne LiDAR Scanning (ALS) and more recently spaceborne LiDAR (GEDI, ICESat-2) '
    'provide direct measurements of canopy height that resolve the spectral ambiguity between '
    'shrubs and trees. Meta\'s Global Canopy Height Model, derived from GEDI calibration of '
    'Sentinel-2 imagery, now provides wall-to-wall canopy height estimates at 1-meter resolution, '
    'making structural information available for any terrestrial location (Tolan et al., 2024).'
)

doc.add_heading('2.2 Machine Learning for Vegetation Classification', level=2)
doc.add_paragraph(
    'Pixel-based classification using ensemble methods (Random Forest, Gradient Boosting) remains '
    'the dominant approach for vegetation mapping from multi-source remote sensing data. These '
    'methods handle heterogeneous feature spaces (spectral, textural, structural) naturally and '
    'require relatively small training datasets compared to deep learning. XGBoost (Chen & '
    'Guestrin, 2016) has shown particular strength in remote sensing applications due to its '
    'regularization capabilities and handling of missing values.'
)
doc.add_paragraph(
    'Deep learning approaches, particularly U-Net (Ronneberger et al., 2015) for semantic '
    'segmentation and Mask R-CNN (He et al., 2017) for instance segmentation, have achieved '
    'state-of-the-art results in general computer vision. However, their application to ecological '
    'remote sensing is limited by: (1) the requirement for large, pixel-accurate labeled datasets; '
    '(2) sensitivity to domain shift between training and deployment sites; and (3) computational '
    'cost that limits scalability. Recent work has shown that for vegetation classification with '
    'limited training data, gradient-boosted trees with engineered features can match or exceed '
    'CNN performance (Fassnacht et al., 2024).'
)

doc.add_heading('2.3 Foundation Models for Annotation', level=2)
doc.add_paragraph(
    'The Segment Anything Model (SAM) represents a paradigm shift in image annotation. Trained on '
    'over 1 billion masks, SAM can segment arbitrary objects from minimal prompts (points, boxes, '
    'or text). In remote sensing, SAM has been applied to building footprint extraction, crop field '
    'delineation, and tree crown segmentation (Osco et al., 2023). This dissertation extends SAM\'s '
    'application to shrub boundary refinement, where point prompts from field survey coordinates '
    'are used to generate precise shrub outlines that conform to the actual vegetation extent '
    'visible in aerial imagery.'
)
doc.add_paragraph(
    'A novel contribution of this work is the dual use of SAM for both positive label refinement '
    '(Pass 1: converting circles to precise masks) and hard negative mining (Pass 2: automatic '
    'mask generation with height-based classification into tree and rock categories). This '
    'approach addresses the well-known challenge of defining what constitutes a "non-shrub" pixel '
    'in an ecologically meaningful way.'
)
doc.add_page_break()

# ---- CHAPTER 3: STUDY AREA ----
doc.add_heading('Chapter 3: Study Area and Data', level=1)

doc.add_heading('3.1 Site Descriptions', level=2)
doc.add_paragraph(
    'Six field sites in California were selected to represent the ecological diversity of '
    'shrub-bearing landscapes across the state. Site selection criteria included: (1) availability '
    'of TLS data from ongoing monitoring programs; (2) presence of shrub species with heights '
    'in the 1-4 meter range; and (3) coverage by recent NAIP acquisitions (2022 flight year).'
)
t = doc.add_table(rows=7, cols=4, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Site', 'Code', 'Ecosystem', 'Scan Dates']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for r, row in enumerate([
    ('Calaveras Big Trees', 'CATCU', 'Mixed conifer forest', 'Oct 2024, May-Jun 2025'),
    ('DL Bliss State Park', 'CAAEU', 'Sierra Nevada montane', 'Multiple'),
    ('Independence Lake', '-', 'Subalpine lake basin', 'Jul-Aug 2024, Aug 2025'),
    ('Pacific Union College', '-', 'N. California woodland', 'Nov 2024'),
    ('Sedgwick Reserve', 'CASBC', 'Coastal sage scrub', 'Sep-Nov 2024, Jun-Jul 2025'),
    ('Shaver Lake', 'CAFKU', 'Sierra foothills', 'Jul 2024'),
], 1):
    for c, v in enumerate(row):
        t.rows[r].cells[c].text = v

doc.add_heading('3.2 Data Sources', level=2)

doc.add_heading('3.2.1 Terrestrial LiDAR Scanning (TLS)', level=3)
doc.add_paragraph(
    'High-resolution TLS point clouds were acquired at each site using tripod-mounted scanners. '
    'These scans provide centimeter-level 3D structural information about individual shrubs, '
    'enabling precise georeferencing of shrub locations. Rigid-body transformation matrices align '
    'TLS scans to the NAIP coordinate system (EPSG:6350, NAD83(2011) / Conus Albers). Field-verified '
    'shrub survey lists provide species, location, and structural attributes for each shrub.'
)

doc.add_heading('3.2.2 NAIP Aerial Imagery', level=3)
doc.add_paragraph(
    'The National Agriculture Imagery Program (NAIP) provides 4-band aerial imagery (Red, Green, '
    'Blue, and Near-Infrared) at 1-meter ground sample distance. NAIP imagery from the 2022 '
    'acquisition cycle was used for all sites to ensure temporal consistency. For the autonomous '
    'agent, NAIP tiles are retrieved programmatically from Microsoft Planetary Computer via the '
    'STAC API.'
)

doc.add_heading('3.2.3 Canopy Height Model (CHM)', level=3)
doc.add_paragraph(
    'Canopy height data is derived from Meta\'s Global Canopy Height Model, accessed through '
    'Google Earth Engine. This dataset provides wall-to-wall canopy height estimates at 1-meter '
    'resolution, calibrated against GEDI spaceborne LiDAR measurements. CHM values are used both '
    'as input features for the classifier and as validation criteria for SAM annotation refinement '
    '(shrub height range: 1.0-4.0 m).'
)
doc.add_page_break()

# Save document to disk for Part 2
out = r"c:\Users\sefak\OneDrive\Documents\sefakarabas\OneDrive\Desktop\Dataset-2\Shrubwise_Thesis.docx"
doc.save(out)
print(f"Part 1 saved to: {out}")
